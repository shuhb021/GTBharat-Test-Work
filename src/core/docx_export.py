"""
FAR Automation Tool — Word Document Export
Generates a professional Word document with cover page, analysis tables,
and AI remarks summary using python-docx.
"""

import os
import logging
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

logger = logging.getLogger(__name__)


def _set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading_elm = cell._element.get_or_add_tcPr()
    shading = shading_elm.makeelement(qn('w:shd'), {
        qn('w:fill'): color_hex,
        qn('w:val'): 'clear'
    })
    shading_elm.append(shading)


def _add_table_row(table, values, is_header=False, highlight=False):
    """Add a row to a Word table with styling."""
    row = table.add_row()
    for i, val in enumerate(values):
        cell = row.cells[i]
        cell.text = str(val) if val is not None else ''
        
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.RIGHT
            for run in paragraph.runs:
                run.font.size = Pt(9)
                run.font.name = 'Segoe UI'
                if is_header:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        
        if is_header:
            _set_cell_shading(cell, '252526')
        elif highlight:
            _set_cell_shading(cell, 'EAB308')
    
    return row


def export_docx(output_path, bs_result, pl_result, ratios,
                    remarks_bs, remarks_pl,
                    client_name, firm_name, financial_year,
                    rounding_unit, cy_year, py_year, chart_paths=None,
                    meta=None):
    """
    Generate the FAR Word document.
    """
    if not HAS_DOCX:
        logger.error("python-docx not installed — cannot generate Word document")
        return False
    
    doc = Document()
    
    # Add footer with generation timestamp
    for section in doc.sections:
        footer = section.footer
        footer_p = footer.paragraphs[0]
        footer_p.text = f"FAR Workpaper — Generated: {datetime.now().strftime('%d %B %Y, %H:%M:%S')}"
        footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Segoe UI'
    font.size = Pt(10)
    
    # ── Cover Page ──
    doc.add_paragraph('')
    
    # Firm Name
    firm = doc.add_paragraph()
    firm.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = firm.add_run(firm_name if firm_name else "Walker Chandiok & Co LLP")
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x70, 0x30, 0xA0) # Purple color
    
    doc.add_paragraph('')
    
    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Final Balance Analysis')
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x70, 0x30, 0xA0)
    
    doc.add_paragraph('')
    
    # Details table-like paragraphs
    p_client = doc.add_paragraph()
    p_client.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_client.add_run(f'Client: {client_name}')
    run.font.size = Pt(12)
    run.font.bold = True
    
    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_date.add_run(f'Analysis Date: {datetime.now().strftime("%B %d, %Y")}')
    run.font.size = Pt(12)
    
    doc.add_paragraph('')
    
    # Year-over-Year Financial Comparison Header
    yoy_hdr = doc.add_paragraph()
    yoy_hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = yoy_hdr.add_run('Year-over-Year Financial Comparison')
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x70, 0x30, 0xA0)
    
    if meta:
        py_days = meta.get('py_days', 365)
        py_months = meta.get('py_months', 12.0)
        py_cov = meta.get('py_coverage', 100.0)
        cy_days = meta.get('cy_days', 365)
        cy_months = meta.get('cy_months', 12.0)
        cy_cov = meta.get('cy_coverage', 100.0)
    else:
        py_days = cy_days = 365
        py_months = cy_months = 12.0
        py_cov = cy_cov = 100.0

    p_py = doc.add_paragraph(style='List Bullet')
    run = p_py.add_run(f"Previous Year ({py_year}): Days: {py_days} | Months: {py_months} | Coverage: {py_cov}%")
    run.font.size = Pt(11)
    
    p_cy = doc.add_paragraph(style='List Bullet')
    run = p_cy.add_run(f"Current Year ({cy_year}): Days: {cy_days} | Months: {cy_months} | Coverage: {cy_cov}%")
    run.font.size = Pt(11)

    doc.add_paragraph('')
    
    # Analysis Summary Header
    summary_hdr = doc.add_paragraph()
    summary_hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = summary_hdr.add_run('Analysis Summary')
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x70, 0x30, 0xA0)
    
    # Summary Bullet Points
    bullets = [
        "Balance Sheet Analysis",
        "Profit and Loss Analysis",
        "Ratio and Relations Analysis",
        "Visual Analytics and Charts",
        "Notes ",
    ]
    for bullet in bullets:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(bullet)
        run.font.size = Pt(11)
        
    doc.add_paragraph('')
    doc.add_paragraph('')
    
    # Confidentiality
    conf = doc.add_paragraph()
    conf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = conf.add_run('Confidential - For Internal Use Only')
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x7F, 0x7F, 0x7F)
    
    doc.add_page_break()
    
    # ── Balance Sheet Analysis ──
    doc.add_heading('Balance Sheet Analysis', level=1)
    doc.add_paragraph(
        f'As at 31 March {cy_year} vs 31 March {py_year}  |  (Rs. in {rounding_unit})')
    
    bs_headers = ['Particulars', 'Note', f'CY ({cy_year})', f'PY ({py_year})',
                  'Variance', 'Var %', 'Remarks']
    bs_table = doc.add_table(rows=1, cols=len(bs_headers))
    bs_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    bs_table.style = 'Table Grid'
    
    # Header row
    for i, h in enumerate(bs_headers):
        cell = bs_table.rows[0].cells[i]
        cell.text = h
        _set_cell_shading(cell, '252526')
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    
def _fmt_val(val, fmt):
    if val is None:
        return '-NA-'
    if val == '-':
        return '-'
    try:
        v = float(val)
        if v == 0:
            return '-'
        return f"{v:{fmt}}"
    except (TypeError, ValueError):
        return str(val)

def _write_docx_section(doc, heading, items):
    if not items:
        return
    for _ in range(2):
        doc.add_paragraph('')
    # Add section subheading
    p = doc.add_paragraph()
    run = p.add_run(heading)
    run.font.name = 'Segoe UI'
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x70, 0x30, 0xA0) # Purple theme
    
    max_cols = max(len(row) for row in items) if items else 2
    if max_cols < 2:
        max_cols = 2
    table = doc.add_table(rows=0, cols=max_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Normal Table'
    col_width = Inches(6.5 / max_cols)
    for f_row in items:
        row = table.add_row()
        for i in range(max_cols):
            cell = row.cells[i]
            cell.width = col_width
            cell_val = f_row[i] if i < len(f_row) else ''
            cell.text = str(cell_val).strip()
            for p_cell in cell.paragraphs:
                p_cell.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run_cell in p_cell.runs:
                    run_cell.font.name = 'Segoe UI'
                    run_cell.font.size = Pt(9)
                    run_cell.font.bold = False


def export_docx(output_path, bs_result, pl_result, ratios,
                remarks_bs, remarks_pl,
                client_name, firm_name, financial_year,
                rounding_unit, cy_year, py_year, chart_paths=None,
                bs_notes=None, bs_signatures=None, bs_footers=None,
                pl_notes=None, pl_signatures=None, pl_footers=None,
                notes_sheet_notes=None, notes_signatures=None, notes_footers=None,
                report_type='FAR'):
    """
    Generate the FAR Word document.
    """
    if not HAS_DOCX:
        logger.error("python-docx not installed — cannot generate Word document")
        return False
    
    doc = Document()
    
    _REPORT_LABELS = {
        'FAR': 'Final Analytical Report',
        'PAR': 'Preliminary Analytical Report',
    }
    report_label = _REPORT_LABELS.get(report_type, 'Final Analytical Report')

    # Add footer with generation timestamp
    for section in doc.sections:
        footer = section.footer
        footer_p = footer.paragraphs[0]
        footer_p.text = f"{report_label} Workpaper — Generated: {datetime.now().strftime('%d %B %Y, %H:%M:%S')}"
        footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Segoe UI'
    font.size = Pt(10)
    
    # ── Cover Page ──
    doc.add_paragraph('')
    
    # Firm Name
    firm = doc.add_paragraph()
    firm.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = firm.add_run(firm_name if firm_name else "Walker Chandiok & Co LLP")
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x70, 0x30, 0xA0) # Purple color
    
    doc.add_paragraph('')
    
    # Subtitle — dynamic report type
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(report_label)
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x70, 0x30, 0xA0)
    
    doc.add_paragraph('')
    
    # Details table-like paragraphs
    p_client = doc.add_paragraph()
    p_client.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_client.add_run(f'Client: {client_name}')
    run.font.size = Pt(12)
    run.font.bold = True
    
    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_date.add_run(f'Analysis Date: {datetime.now().strftime("%B %d, %Y")}')
    run.font.size = Pt(12)
    
    doc.add_paragraph('')
    
    # Analysis Summary Header
    summary_hdr = doc.add_paragraph()
    summary_hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = summary_hdr.add_run('Analysis Summary')
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x70, 0x30, 0xA0)
    
    # Summary Bullet Points
    bullets = [
        "Balance Sheet Analysis",
        "Profit and Loss Analysis",
        "Ratio and Relations Analysis",
        "Visual Analytics and Charts",
        "Notes ",
    ]
    for bullet in bullets:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(bullet)
        run.font.size = Pt(11)
        
    doc.add_paragraph('')
    doc.add_paragraph('')
    
    # Confidentiality
    conf = doc.add_paragraph()
    conf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = conf.add_run('Confidential - For Internal Use Only')
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x7F, 0x7F, 0x7F)
    
    doc.add_page_break()
    
    # ── Balance Sheet Analysis ──
    doc.add_heading('Balance Sheet Analysis', level=1)
    doc.add_paragraph(
        f'As at 31 March {cy_year} vs 31 March {py_year}  |  (Rs. in {rounding_unit})')
    
    bs_headers = ['Particulars', 'Note', f'CY ({cy_year})', f'PY ({py_year})',
                  'Variance', 'Var %', 'Remarks']
    bs_table = doc.add_table(rows=1, cols=len(bs_headers))
    bs_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    bs_table.style = 'Table Grid'
    
    # Header row
    for i, h in enumerate(bs_headers):
        cell = bs_table.rows[0].cells[i]
        cell.text = h
        _set_cell_shading(cell, '252526')
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    
    fmt = ",.2f" if rounding_unit in ('Lakhs', 'Millions', 'Actuals') else ",.0f"
    for idx, row_data in enumerate(bs_result):
        remark = remarks_bs.get(idx, '')
        values = [
            row_data.get('particulars', ''),
            row_data.get('note', ''),
            _fmt_val(row_data.get('cy'), fmt),
            _fmt_val(row_data.get('py'), fmt),
            _fmt_val(row_data.get('variance_abs'), fmt),
            row_data.get('display_pct', ''),
            remark
        ]
        _add_table_row(bs_table, values,
                      highlight=row_data.get('flag', False))
    
    _write_docx_section(doc, "Explanatory Notes & Comments", bs_notes)
    _write_docx_section(doc, "Signatures & Sign-offs", bs_signatures)
    _write_docx_section(doc, "Footers & Other Metadata", bs_footers)
    doc.add_page_break()
    
    # ── Profit & Loss Analysis ──
    doc.add_heading('Profit & Loss Analysis', level=1)
    doc.add_paragraph(
        f'For the year ended 31 March {cy_year} vs 31 March {py_year}  |  (Rs. in {rounding_unit})')
    
    pl_table = doc.add_table(rows=1, cols=len(bs_headers))
    pl_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    pl_table.style = 'Table Grid'
    
    for i, h in enumerate(bs_headers):
        cell = pl_table.rows[0].cells[i]
        cell.text = h
        _set_cell_shading(cell, '252526')
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    
    for idx, row_data in enumerate(pl_result):
        remark = remarks_pl.get(idx, '')
        values = [
            row_data.get('particulars', ''),
            row_data.get('note', ''),
            _fmt_val(row_data.get('cy'), fmt),
            _fmt_val(row_data.get('py'), fmt),
            _fmt_val(row_data.get('variance_abs'), fmt),
            row_data.get('display_pct', ''),
            remark
        ]
        _add_table_row(pl_table, values,
                      highlight=row_data.get('flag', False))
    
    _write_docx_section(doc, "Explanatory Notes & Comments", pl_notes)
    _write_docx_section(doc, "Signatures & Sign-offs", pl_signatures)
    _write_docx_section(doc, "Footers & Other Metadata", pl_footers)
    doc.add_page_break()
    
    # ── Ratio Analysis ──
    doc.add_heading('Ratio Analysis', level=1)
    doc.add_paragraph(
        f'As at 31 March {cy_year} vs 31 March {py_year}')
    
    ratio_headers = ['Ratio', 'Formula', 'CY', 'PY', 'Change %']
    ratio_table = doc.add_table(rows=1, cols=len(ratio_headers))
    ratio_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    ratio_table.style = 'Table Grid'
    
    for i, h in enumerate(ratio_headers):
        cell = ratio_table.rows[0].cells[i]
        cell.text = h
        _set_cell_shading(cell, '252526')
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    
    for ratio in ratios:
        values = [
            ratio['key'],
            ratio['formula'],
            _fmt_val(ratio.get('cy'), ".2f"),
            _fmt_val(ratio.get('py'), ".2f"),
            ratio.get('display_change', '')
        ]
        _add_table_row(ratio_table, values,
                      highlight=ratio.get('flag', False))
    
    # ── AI Remarks Summary ──
    if remarks_bs or remarks_pl:
        doc.add_page_break()
        doc.add_heading('AI-Generated Audit Remarks Summary', level=1)
        
        if remarks_bs:
            doc.add_heading('Balance Sheet Remarks', level=2)
            for idx, remark in sorted(remarks_bs.items()):
                if idx < len(bs_result):
                    item_name = bs_result[idx].get('particulars', f'Item {idx}')
                    p = doc.add_paragraph()
                    run = p.add_run(f'{item_name}: ')
                    run.font.bold = True
                    run.font.size = Pt(10)
                    run = p.add_run(remark)
                    run.font.size = Pt(10)
        
        if remarks_pl:
            doc.add_heading('Profit & Loss Remarks', level=2)
            for idx, remark in sorted(remarks_pl.items()):
                if idx < len(pl_result):
                    item_name = pl_result[idx].get('particulars', f'Item {idx}')
                    p = doc.add_paragraph()
                    run = p.add_run(f'{item_name}: ')
                    run.font.bold = True
                    run.font.size = Pt(10)
                    run = p.add_run(remark)
                    run.font.size = Pt(10)
    
    # ── Visual Analytics Dashboard ──
    if chart_paths:
        doc.add_page_break()
        doc.add_heading('Visual Analytics Dashboard', level=1)
        doc.add_paragraph('Selected key visualizations and composition analyses.')

        for idx, path in enumerate(chart_paths):
            if os.path.exists(path):
                try:
                    chart_titles = [
                        'Revenue vs Expenses (CY/PY)',
                        'Profit Before Tax Trend',
                        'Asset Composition (CY)',
                        'Liability & Equity Composition (CY)',
                        'Top 5 BS Variances',
                        'Top 5 PL Variances'
                    ]
                    if idx < len(chart_titles):
                        doc.add_heading(chart_titles[idx], level=2)
                    
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run()
                    run.add_picture(path, width=Inches(5.0))
                    
                    doc.add_paragraph('')
                except Exception as e:
                    logger.error("Failed to add chart %d to Word document: %s", idx, e)
    doc.save(output_path)
    logger.info("FAR Word document saved to %s", output_path)
    return True
